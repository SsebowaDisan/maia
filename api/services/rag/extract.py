"""RAG Pipeline Phase 3: Extract — pull text + evidence geometry from sources."""

from __future__ import annotations

import re
from html.parser import HTMLParser

from api.services.rag.types import (
    BoundingBox,
    ClassificationResult,
    EvidenceSpan,
    ExtractionResult,
    FigureSpan,
    FormulaSpan,
    HeadingSpan,
    PageExtraction,
    ProcessingRoute,
    SourceRecord,
    SourceType,
    TableSpan,
)

# ── LaTeX / math detection patterns ────────────────────────────────────────

_LATEX_BLOCK = re.compile(
    r"(\$\$.*?\$\$|\\\[.*?\\\]|\\begin\{(?:equation|align|gather)\}.*?"
    r"\\end\{(?:equation|align|gather)\})",
    re.DOTALL,
)

_LATEX_INLINE = re.compile(r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)")

_MATH_UNICODE_RE = re.compile(
    r"[\u2200-\u22FF\u2A00-\u2AFF\u00B1\u00D7\u00F7\u2260\u2264\u2265\u221E]+"
)

# ── Heading detection helpers ───────────────────────────────────────────────

_HEADING_PATTERNS = re.compile(
    r"^(?:"
    r"(?:Chapter|CHAPTER)\s+\d+|"              # Chapter 1
    r"\d+(?:\.\d+)*\s+[A-Z]|"                  # 3.2 Results
    r"[A-Z][A-Z\s]{4,}$|"                      # ALL CAPS LINE
    r"#{1,6}\s+"                                # Markdown headings
    r")",
    re.MULTILINE,
)


def _guess_heading_level(text: str) -> int:
    """Heuristic heading level: 1 for 'Chapter', 2 for numbered, 3 for caps."""
    stripped = text.strip()
    if stripped.startswith("#"):
        return stripped.split()[0].count("#")
    if stripped.upper().startswith("CHAPTER"):
        return 1
    if re.match(r"^\d+\s+", stripped):
        return 2
    if re.match(r"^\d+\.\d+", stripped):
        return 3
    if stripped.isupper() and len(stripped) > 4:
        return 2
    return 3


# ── Image extraction + vision description ──────────────────────────────────

import base64
import logging
import os

import httpx

_logger = logging.getLogger(__name__)


async def _extract_images_from_pdf(file_data: bytes, source_id: str) -> list[FigureSpan]:
    """Extract images from PDF pages and describe them using a vision LLM.
    Makes diagrams, charts, figures, and photos searchable via text embeddings.
    """
    figures: list[FigureSpan] = []

    try:
        import fitz
    except ImportError:
        return figures

    doc = fitz.open(stream=file_data, filetype="pdf")

    for page_idx in range(doc.page_count):
        page = doc[page_idx]
        image_list = page.get_images(full=True)

        for img_idx, img_info in enumerate(image_list):
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                if not base_image:
                    continue

                image_bytes = base_image["image"]
                image_ext = base_image.get("ext", "png")
                width = base_image.get("width", 0)
                height = base_image.get("height", 0)

                # Skip tiny images (icons, bullets, etc)
                if width < 50 or height < 50:
                    continue

                # Convert to base64
                b64 = base64.b64encode(image_bytes).decode("utf-8")
                mime = f"image/{image_ext}" if image_ext != "jpg" else "image/jpeg"

                # Get image position on page
                bbox = None
                for img_rect in page.get_image_rects(xref):
                    bbox = BoundingBox(
                        x=img_rect.x0, y=img_rect.y0,
                        width=img_rect.x1 - img_rect.x0,
                        height=img_rect.y1 - img_rect.y0,
                        page=page_idx,
                    )
                    break

                # Describe the image using vision LLM
                description = await _describe_image(b64, mime, page_idx, source_id)

                if description:
                    char_start = 0  # Will be set during normalization
                    figure = FigureSpan(
                        text=f"[Figure on page {page_idx + 1}]: {description}",
                        source_id=source_id,
                        page=page_idx,
                        char_start=char_start,
                        char_end=char_start + len(description),
                        bbox=bbox,
                        image_data=b64[:100] + "...",  # Store truncated for reference
                        caption=description,
                        alt_text=description,
                    )
                    figures.append(figure)

            except Exception as exc:
                _logger.debug("Failed to extract image %d from page %d: %s", img_idx, page_idx, exc)

    doc.close()
    return figures


async def _describe_image(b64_data: str, mime_type: str, page: int, source_id: str, model: str = "") -> str:
    """Send an image to vision model to get a text description."""
    api_key = os.environ.get("MAIA_RAG_API_KEY", "") or os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        return ""

    base_url = os.environ.get("MAIA_RAG_LLM_BASE_URL", "https://api.openai.com/v1")

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"{base_url}/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": model or os.environ.get("MAIA_RAG_VISION_MODEL", "gpt-4o-mini"),
                    "messages": [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": (
                                        "Describe this figure/diagram/chart from a technical PDF. "
                                        "Include: what it shows, any labels, axes, values, units, "
                                        "trends, and key data points. Be specific and factual. "
                                        "If it contains equations or formulas, write them out. "
                                        "If it's a data table, list the key values."
                                    ),
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:{mime_type};base64,{b64_data}",
                                        "detail": "high",
                                    },
                                },
                            ],
                        }
                    ],
                    "max_tokens": 500,
                },
            )

            if response.status_code != 200:
                _logger.debug("Vision API returned %d for image on page %d", response.status_code, page)
                return ""

            data = response.json()
            return data["choices"][0]["message"]["content"].strip()

    except Exception as exc:
        _logger.debug("Vision description failed for page %d: %s", page, exc)
        return ""


async def _enrich_with_image_descriptions(
    result: ExtractionResult, file_data: bytes, source_id: str,
) -> ExtractionResult:
    """Post-process: extract images and add vision descriptions to the extraction result.
    Works with any extraction method (Docling, PyMuPDF, etc).
    """
    figures = await _extract_images_from_pdf(file_data, source_id)

    if not figures:
        return result

    _logger.info("Extracted %d image descriptions from %s", len(figures), source_id)

    # Add figure descriptions to the extraction
    result.figures.extend(figures)

    # Also add figure text to page extractions and full text
    for fig in figures:
        # Find or create the page
        page_found = False
        for page_ext in result.pages:
            if page_ext.page_number == fig.page:
                page_ext.figures.append(fig)
                page_ext.text += f"\n\n{fig.text}"
                page_ext.spans.append(EvidenceSpan(
                    text=fig.text,
                    source_id=source_id,
                    page=fig.page,
                    char_start=len(result.full_text),
                    char_end=len(result.full_text) + len(fig.text),
                    bbox=fig.bbox,
                ))
                page_found = True
                break

        if not page_found:
            result.pages.append(PageExtraction(
                page_number=fig.page,
                text=fig.text,
                figures=[fig],
                spans=[EvidenceSpan(
                    text=fig.text,
                    source_id=source_id,
                    page=fig.page,
                    char_start=len(result.full_text),
                    char_end=len(result.full_text) + len(fig.text),
                    bbox=fig.bbox,
                )],
                char_offset=len(result.full_text),
            ))

        result.full_text += f"\n\n{fig.text}"

    return result


# ── PDF native text extraction (PyMuPDF) ───────────────────────────────────


def _extract_pdf_native(file_data: bytes, source_id: str) -> ExtractionResult:
    """Extract text from a native-text PDF with per-block bounding boxes."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_data, filetype="pdf")
    pages: list[PageExtraction] = []
    all_formulas: list[FormulaSpan] = []
    all_tables: list[TableSpan] = []
    full_parts: list[str] = []
    char_cursor = 0
    warnings: list[str] = []

    for page_idx in range(doc.page_count):
        page = doc[page_idx]
        page_rect = page.rect  # (x0, y0, x1, y1) in PDF points
        page_width = page_rect.width or 595   # fallback to A4
        page_height = page_rect.height or 842
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

        page_text_parts: list[str] = []
        page_spans: list[EvidenceSpan] = []
        page_headings: list[HeadingSpan] = []
        page_formulas: list[FormulaSpan] = []
        page_tables: list[TableSpan] = []
        page_char_offset = char_cursor

        for block in blocks:
            if block["type"] != 0:  # skip image blocks
                continue

            block_text_parts: list[str] = []
            for line in block.get("lines", []):
                line_text = ""
                for span in line.get("spans", []):
                    line_text += span["text"]
                block_text_parts.append(line_text)

            block_text = "\n".join(block_text_parts).strip()
            if not block_text:
                continue

            bbox_raw = block["bbox"]  # (x0, y0, x1, y1)
            bbox = BoundingBox(
                x=bbox_raw[0],
                y=bbox_raw[1],
                width=bbox_raw[2] - bbox_raw[0],
                height=bbox_raw[3] - bbox_raw[1],
                page=page_idx,
                page_width=page_width,
                page_height=page_height,
            )

            span_start = char_cursor
            span_end = char_cursor + len(block_text)

            ev_span = EvidenceSpan(
                text=block_text,
                source_id=source_id,
                page=page_idx,
                char_start=span_start,
                char_end=span_end,
                bbox=bbox,
            )
            page_spans.append(ev_span)

            # Detect headings by font size — blocks where the largest span
            # font size exceeds the median are heading candidates.
            font_sizes: list[float] = []
            for line in block.get("lines", []):
                for sp in line.get("spans", []):
                    font_sizes.append(sp.get("size", 12.0))
            max_font = max(font_sizes) if font_sizes else 12.0

            if max_font >= 14.0 and len(block_text) < 200:
                page_headings.append(
                    HeadingSpan(
                        text=block_text,
                        level=1 if max_font >= 20 else (2 if max_font >= 16 else 3),
                        page=page_idx,
                        char_start=span_start,
                        char_end=span_end,
                    )
                )

            # Detect formulas in block text
            for m in _LATEX_BLOCK.finditer(block_text):
                formula_start = span_start + m.start()
                formula_end = span_start + m.end()
                page_formulas.append(
                    FormulaSpan(
                        text=m.group(),
                        source_id=source_id,
                        page=page_idx,
                        char_start=formula_start,
                        char_end=formula_end,
                        bbox=bbox,
                        latex=m.group(),
                    )
                )

            page_text_parts.append(block_text)
            char_cursor = span_end + 1  # +1 for newline separator

        page_text = "\n".join(page_text_parts)
        full_parts.append(page_text)

        pages.append(
            PageExtraction(
                page_number=page_idx,
                text=page_text,
                headings=page_headings,
                spans=page_spans,
                formulas=page_formulas,
                tables=page_tables,
                char_offset=page_char_offset,
            )
        )
        all_formulas.extend(page_formulas)
        all_tables.extend(page_tables)

    doc.close()
    full_text = "\n".join(full_parts)

    return ExtractionResult(
        source_id=source_id,
        pages=pages,
        formulas=all_formulas,
        tables=all_tables,
        full_text=full_text,
        extraction_method="native_text",
        warnings=warnings,
    )


# ── OCR placeholder ────────────────────────────────────────────────────────


def _extract_pdf_ocr(
    file_data: bytes, source_id: str, scientific: bool = False,
) -> ExtractionResult:
    """Extract text from scanned/image PDFs using Docling.

    Docling handles: OCR, layout analysis, tables, formulas, figures, reading order.
    Falls back to PyMuPDF native text if Docling is not installed.
    """
    # Try Docling first
    try:
        return _extract_with_docling(file_data, source_id)
    except ImportError:
        import logging
        logging.getLogger(__name__).warning(
            "Docling not installed — falling back to PyMuPDF native text. "
            "Install with: pip install docling"
        )
    except Exception as exc:
        import logging
        logging.getLogger(__name__).warning("Docling extraction failed: %s — falling back to PyMuPDF", exc)

    # Fallback: PyMuPDF native text (works for PDFs with embedded text layer)
    return _extract_pdf_native_fallback(file_data, source_id)


def _extract_with_docling(file_data: bytes, source_id: str) -> ExtractionResult:
    """Use Docling for full document understanding — OCR, tables, formulas, layout."""
    import tempfile
    import os

    # Fix Windows symlink issue with HuggingFace cache
    os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS_WARNING", "1")
    os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS", "1")

    from docling.document_converter import DocumentConverter

    # Docling needs a file path, not bytes
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(file_data)
        tmp_path = tmp.name

    try:
        converter = DocumentConverter()
        result = converter.convert(tmp_path)
        doc = result.document

        pages: list[PageExtraction] = {}  # type: ignore
        page_map: dict[int, PageExtraction] = {}
        all_formulas: list[FormulaSpan] = []
        all_tables: list[TableSpan] = []
        full_parts: list[str] = []
        char_cursor = 0

        # Process document elements — Docling provides structured output
        for item, _level in doc.iterate_items():
            item_text = item.text if hasattr(item, "text") else str(item)
            if not item_text or not item_text.strip():
                continue

            # Determine page number
            page_num = 0
            if hasattr(item, "prov") and item.prov:
                prov = item.prov[0] if isinstance(item.prov, list) else item.prov
                page_num = getattr(prov, "page_no", 1) - 1  # 0-indexed
                if page_num < 0:
                    page_num = 0

            # Get bounding box if available
            bbox = None
            if hasattr(item, "prov") and item.prov:
                prov = item.prov[0] if isinstance(item.prov, list) else item.prov
                if hasattr(prov, "bbox"):
                    b = prov.bbox
                    if hasattr(b, "l"):  # Docling BoundingBox has l, t, r, b
                        bbox = BoundingBox(
                            x=float(b.l), y=float(b.t),
                            width=float(b.r - b.l), height=float(b.b - b.t),
                            page=page_num,
                        )
                    elif hasattr(b, "x"):
                        bbox = BoundingBox(
                            x=float(b.x), y=float(b.y),
                            width=float(getattr(b, "width", 0)),
                            height=float(getattr(b, "height", 0)),
                            page=page_num,
                        )

            # Create evidence span
            span_start = char_cursor
            span_end = char_cursor + len(item_text)
            ev_span = EvidenceSpan(
                text=item_text,
                source_id=source_id,
                page=page_num,
                char_start=span_start,
                char_end=span_end,
                bbox=bbox,
            )

            # Ensure page exists
            if page_num not in page_map:
                page_map[page_num] = PageExtraction(
                    page_number=page_num,
                    text="",
                    spans=[],
                    headings=[],
                    formulas=[],
                    tables=[],
                    char_offset=char_cursor,
                )

            page_ext = page_map[page_num]
            page_ext.spans.append(ev_span)

            # Detect element type
            label = getattr(item, "label", "") or ""
            label_lower = label.lower() if isinstance(label, str) else ""

            # Headings
            if "heading" in label_lower or "title" in label_lower or "section" in label_lower:
                level = 1 if "title" in label_lower else (2 if "section" in label_lower else 3)
                page_ext.headings.append(HeadingSpan(
                    text=item_text, level=level, page=page_num,
                    char_start=span_start, char_end=span_end,
                ))

            # Formulas / equations
            if "formula" in label_lower or "equation" in label_lower or "math" in label_lower:
                formula = FormulaSpan(
                    text=item_text, source_id=source_id, page=page_num,
                    char_start=span_start, char_end=span_end, bbox=bbox,
                    latex=item_text,
                )
                page_ext.formulas.append(formula)
                all_formulas.append(formula)

            # Tables
            if "table" in label_lower:
                rows = []
                headers = []
                if hasattr(item, "export_to_dataframe"):
                    try:
                        df = item.export_to_dataframe()
                        headers = list(df.columns)
                        rows = df.values.tolist()
                    except Exception:
                        pass
                table = TableSpan(
                    text=item_text, source_id=source_id, page=page_num,
                    char_start=span_start, char_end=span_end, bbox=bbox,
                    rows=[[str(c) for c in r] for r in rows],
                    headers=[str(h) for h in headers],
                )
                page_ext.tables.append(table)
                all_tables.append(table)

            page_ext.text += item_text + "\n"
            full_parts.append(item_text)
            char_cursor = span_end + 1

        # Sort pages
        sorted_pages = [page_map[k] for k in sorted(page_map.keys())]
        full_text = "\n".join(full_parts)

        return ExtractionResult(
            source_id=source_id,
            pages=sorted_pages,
            formulas=all_formulas,
            tables=all_tables,
            full_text=full_text,
            extraction_method="docling",
        )
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


def _extract_pdf_native_fallback(file_data: bytes, source_id: str) -> ExtractionResult:
    """Fallback: extract whatever text PyMuPDF can find from a scanned PDF.
    Usually gets the embedded text layer if it exists.
    """
    try:
        import fitz
        doc = fitz.open(stream=file_data, filetype="pdf")
        pages: list[PageExtraction] = []
        full_parts: list[str] = []
        char_cursor = 0

        for page_idx in range(doc.page_count):
            page_text = doc[page_idx].get_text("text").strip()
            if not page_text:
                continue

            span = EvidenceSpan(
                text=page_text, source_id=source_id, page=page_idx,
                char_start=char_cursor, char_end=char_cursor + len(page_text),
            )
            pages.append(PageExtraction(
                page_number=page_idx, text=page_text,
                spans=[span], char_offset=char_cursor,
            ))
            full_parts.append(page_text)
            char_cursor += len(page_text) + 1

        doc.close()
        full_text = "\n".join(full_parts)

        if not full_text.strip():
            return ExtractionResult(
                source_id=source_id,
                warnings=["No text extracted — install Docling for scanned PDF support: pip install docling"],
                extraction_method="pymupdf_fallback",
            )

        return ExtractionResult(
            source_id=source_id, pages=pages, full_text=full_text,
            extraction_method="pymupdf_fallback",
        )
    except Exception as exc:
        return ExtractionResult(
            source_id=source_id,
            warnings=[f"Extraction failed: {exc}. Install Docling: pip install docling"],
            extraction_method="failed",
        )


# ── HTML extraction ────────────────────────────────────────────────────────


class _HTMLTextExtractor(HTMLParser):
    """Simple HTML→text extractor that tracks paragraph positions and headings."""

    def __init__(self) -> None:
        super().__init__()
        self._result: list[str] = []
        self._current: list[str] = []
        self._tag_stack: list[str] = []
        self.headings: list[HeadingSpan] = []
        self.paragraph_spans: list[EvidenceSpan] = []
        self._char_cursor = 0
        self._source_id = ""
        self._skip_tags = {"script", "style", "noscript"}
        self._heading_tags = {"h1", "h2", "h3", "h4", "h5", "h6"}

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        self._tag_stack.append(tag)

    def handle_endtag(self, tag: str) -> None:
        # Flush accumulated text on block-level close
        if tag in ("p", "div", "li", "td", "th", "blockquote", "pre") or tag in self._heading_tags:
            text = "".join(self._current).strip()
            if text:
                start = self._char_cursor
                end = start + len(text)

                self.paragraph_spans.append(
                    EvidenceSpan(
                        text=text,
                        source_id=self._source_id,
                        page=0,
                        char_start=start,
                        char_end=end,
                        paragraph_index=len(self.paragraph_spans),
                    )
                )

                if tag in self._heading_tags:
                    level = int(tag[1])
                    self.headings.append(
                        HeadingSpan(text=text, level=level, page=0, char_start=start, char_end=end)
                    )

                self._result.append(text)
                self._char_cursor = end + 1  # newline separator
            self._current = []

        if self._tag_stack and self._tag_stack[-1] == tag:
            self._tag_stack.pop()

    def handle_data(self, data: str) -> None:
        if any(t in self._skip_tags for t in self._tag_stack):
            return
        self._current.append(data)

    def get_text(self) -> str:
        # Flush remainder
        remainder = "".join(self._current).strip()
        if remainder:
            start = self._char_cursor
            end = start + len(remainder)
            self.paragraph_spans.append(
                EvidenceSpan(
                    text=remainder,
                    source_id=self._source_id,
                    page=0,
                    char_start=start,
                    char_end=end,
                    paragraph_index=len(self.paragraph_spans),
                )
            )
            self._result.append(remainder)
            self._current = []
        return "\n".join(self._result)


async def _extract_html(url: str, source_id: str) -> ExtractionResult:
    """Fetch a URL and extract text with paragraph positions."""
    import httpx

    async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:
        resp = await client.get(url)
        resp.raise_for_status()
        html = resp.text

    parser = _HTMLTextExtractor()
    parser._source_id = source_id
    parser.feed(html)
    full_text = parser.get_text()

    page = PageExtraction(
        page_number=0,
        text=full_text,
        headings=parser.headings,
        spans=parser.paragraph_spans,
        char_offset=0,
    )

    return ExtractionResult(
        source_id=source_id,
        pages=[page],
        full_text=full_text,
        extraction_method="html",
    )


# ── Plaintext extraction ───────────────────────────────────────────────────


def _extract_plaintext(text: str, source_id: str) -> ExtractionResult:
    """Extract text from a plain-text file with paragraph-level spans."""
    paragraphs = re.split(r"\n{2,}", text)
    spans: list[EvidenceSpan] = []
    headings: list[HeadingSpan] = []
    char_cursor = 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        start = char_cursor
        end = start + len(para)

        spans.append(
            EvidenceSpan(
                text=para,
                source_id=source_id,
                page=0,
                char_start=start,
                char_end=end,
                paragraph_index=len(spans),
            )
        )

        # Markdown-style headings
        if para.startswith("#"):
            level = _guess_heading_level(para)
            headings.append(HeadingSpan(text=para.lstrip("# "), level=level, page=0, char_start=start, char_end=end))
        elif _HEADING_PATTERNS.match(para):
            level = _guess_heading_level(para)
            headings.append(HeadingSpan(text=para, level=level, page=0, char_start=start, char_end=end))

        char_cursor = end + 1

    full_text = "\n".join(s.text for s in spans)

    page = PageExtraction(
        page_number=0,
        text=full_text,
        headings=headings,
        spans=spans,
        char_offset=0,
    )

    return ExtractionResult(
        source_id=source_id,
        pages=[page],
        full_text=full_text,
        extraction_method="plaintext",
    )


# ── Office doc placeholder ─────────────────────────────────────────────────


def _extract_office(file_data: bytes, source_id: str, filename: str) -> ExtractionResult:
    """Extract text from Office documents.

    Uses python-docx for .docx; other formats return a placeholder.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in ("docx", "doc"):
        try:
            import io
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(file_data))
            paragraphs: list[str] = []
            spans: list[EvidenceSpan] = []
            headings: list[HeadingSpan] = []
            char_cursor = 0

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                start = char_cursor
                end = start + len(text)

                spans.append(
                    EvidenceSpan(
                        text=text,
                        source_id=source_id,
                        page=0,
                        char_start=start,
                        char_end=end,
                        paragraph_index=len(spans),
                    )
                )

                # Detect headings from style
                style_name = (para.style.name or "").lower()
                if "heading" in style_name:
                    try:
                        level = int(style_name.replace("heading", "").strip())
                    except ValueError:
                        level = 2
                    headings.append(
                        HeadingSpan(text=text, level=level, page=0, char_start=start, char_end=end)
                    )

                paragraphs.append(text)
                char_cursor = end + 1

            full_text = "\n".join(paragraphs)
            page = PageExtraction(
                page_number=0,
                text=full_text,
                headings=headings,
                spans=spans,
                char_offset=0,
            )
            return ExtractionResult(
                source_id=source_id,
                pages=[page],
                full_text=full_text,
                extraction_method="python_docx",
            )
        except ImportError:
            pass  # fall through to placeholder

    # Placeholder for xlsx / pptx / unavailable docx
    notice = f"Office extraction for .{ext} not fully wired — install python-docx / openpyxl / python-pptx."
    return ExtractionResult(
        source_id=source_id,
        pages=[],
        full_text="",
        extraction_method="office_placeholder",
        warnings=[notice],
    )


# ── Image OCR placeholder ──────────────────────────────────────────────────


def _extract_image(file_data: bytes, source_id: str) -> ExtractionResult:
    """Placeholder image OCR — in production this calls PaddleOCR or Tesseract."""
    notice = "Image OCR not yet wired — install PaddleOCR or Tesseract backend."
    return ExtractionResult(
        source_id=source_id,
        pages=[],
        full_text="",
        extraction_method="image_ocr_placeholder",
        warnings=[notice],
    )


# ── Public API ──────────────────────────────────────────────────────────────


async def extract_source(
    source: SourceRecord,
    classification: ClassificationResult,
    file_data: bytes | None = None,
) -> ExtractionResult:
    """Route to the correct extractor based on ProcessingRoute.

    Every piece of extracted text carries its position (page, char offset,
    bounding box) so downstream citation highlighting works.
    """
    route = classification.processing_route

    if route == ProcessingRoute.TEXT_NATIVE:
        if file_data is None:
            return ExtractionResult(
                source_id=source.id, warnings=["No file data for native text extraction"]
            )
        result = _extract_pdf_native(file_data, source.id)

    elif route == ProcessingRoute.OCR_STANDARD:
        if file_data is None:
            return ExtractionResult(
                source_id=source.id, warnings=["No file data for OCR extraction"]
            )
        result = _extract_pdf_ocr(file_data, source.id, scientific=False)

    elif route == ProcessingRoute.OCR_SCIENTIFIC:
        if file_data is None:
            return ExtractionResult(
                source_id=source.id, warnings=["No file data for scientific OCR extraction"]
            )
        result = _extract_pdf_ocr(file_data, source.id, scientific=True)

    elif route == ProcessingRoute.HTML_STANDARD:
        result = await _extract_html(source.upload_url, source.id)

    elif route == ProcessingRoute.OFFICE_STANDARD:
        if file_data is None:
            return ExtractionResult(
                source_id=source.id, warnings=["No file data for Office extraction"]
            )
        result = _extract_office(file_data, source.id, source.filename)

    elif route == ProcessingRoute.IMAGE_OCR:
        if file_data is None:
            return ExtractionResult(
                source_id=source.id, warnings=["No file data for image OCR"]
            )
        result = _extract_image(file_data, source.id)

    elif route == ProcessingRoute.PLAINTEXT:
        if file_data is None:
            return ExtractionResult(
                source_id=source.id, warnings=["No file data for plaintext extraction"]
            )
        text = file_data.decode("utf-8", errors="replace")
        result = _extract_plaintext(text, source.id)
    else:
        result = ExtractionResult(
            source_id=source.id,
            warnings=[f"Unknown processing route: {route}"],
        )

    # ── Enrich PDFs with image descriptions ──────────────────────────────
    # For any PDF route: extract images and describe them with vision LLM.
    # This makes diagrams, charts, figures searchable via text embeddings.
    if file_data and source.source_type in (SourceType.PDF,) and route in (
        ProcessingRoute.TEXT_NATIVE,
        ProcessingRoute.OCR_STANDARD,
        ProcessingRoute.OCR_SCIENTIFIC,
    ):
        try:
            result = await _enrich_with_image_descriptions(result, file_data, source.id)
        except Exception as exc:
            _logger.warning("Image enrichment failed (non-fatal): %s", exc)
            result.warnings.append(f"Image description failed: {exc}")

    return result
